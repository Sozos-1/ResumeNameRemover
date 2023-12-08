import streamlit as st
import spacy
import docx
import requests
from io import BytesIO
import re
from pdf2docx import Converter

def convert_pdf_to_docx(input_pdf_stream):
    output_docx_stream = BytesIO()
    with Converter(input_pdf_stream) as converter:
        converter.convert(output_docx_stream)
    output_docx_stream.seek(0)
    return output_docx_stream

def download_file_from_url(url):
    response = requests.get(url)
    if response.status_code == 200:
        return BytesIO(response.content)
    else:
        raise Exception("Failed to download the file from the provided URL.")

def anonymize_document(file_stream, file_name):
    try:
        st.write(f"Processing file: {file_name}")
        nlp = spacy.load("en_core_web_sm")

        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        phone_pattern = r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b'

        if file_name.endswith('.pdf'):
            converted_docx_stream = convert_pdf_to_docx(file_stream)
            doc = docx.Document(converted_docx_stream)
        else:
            doc = docx.Document(BytesIO(file_stream.getvalue()))
        
        anonymized_doc = docx.Document()

        for para in doc.paragraphs:
            anonymized_text = para.text
            
            anonymized_text = re.sub(email_pattern, "[EMAIL REMOVED]", anonymized_text)
            anonymized_text = re.sub(phone_pattern, "[PHONE NUMBER REMOVED]", anonymized_text)

            processed_text = nlp(anonymized_text)
            for entity in processed_text.ents:
                if entity.label_ == "PERSON":
                    anonymized_text = anonymized_text.replace(entity.text, "[ANONYMIZED]")
            
            anonymized_doc.add_paragraph(anonymized_text)

        anonymized_stream = BytesIO()
        anonymized_doc.save(anonymized_stream)
        anonymized_stream.seek(0)

        st.write("Anonymized document is ready for download.")

        st.download_button(label='Download Anonymized Document',
                           data=anonymized_stream,
                           file_name='anonymized_document.docx',
                           mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        st.error(f"An error occurred: {e}")

# Streamlit UI
st.title("Document Anonymizer")

# File uploader
uploaded_file = st.file_uploader("Upload a document", type=["docx", "pdf"])

# URL input
file_url = st.text_input("...or enter a file URL:")

if uploaded_file is not None:
    anonymize_document(uploaded_file, uploaded_file.name)
elif file_url:
    try:
        file_stream = download_file_from_url(file_url)
        file_name = file_url.split("/")[-1]  # Extracting the file name from the URL
        anonymize_document(file_stream, file_name)
    except Exception as e:
        st.error(f"An error occurred while downloading the file: {e}")
