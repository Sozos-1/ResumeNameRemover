import streamlit as st
import spacy
import docx
from io import BytesIO
import re  # Regular expression library

def anonymize_document(uploaded_file):
    try:
        st.write(f"Processing file: {uploaded_file.name}")
        nlp = spacy.load("en_core_web_sm")

        # Regular expression pattern for detecting emails
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'

        # Read the document
        doc = docx.Document(BytesIO(uploaded_file.getvalue()))
        anonymized_doc = docx.Document()

        for para in doc.paragraphs:
            anonymized_text = para.text
            
            # Remove email addresses
            anonymized_text = re.sub(email_pattern, "[EMAIL REMOVED]", anonymized_text)

            processed_text = nlp(anonymized_text)
            for entity in processed_text.ents:
                if entity.label_ == "PERSON":
                    anonymized_text = anonymized_text.replace(entity.text, "[ANONYMIZED]")
            
            anonymized_doc.add_paragraph(anonymized_text)

        # Save the anonymized document
        anonymized_stream = BytesIO()
        anonymized_doc.save(anonymized_stream)
        anonymized_stream.seek(0)

        st.write("Anonymized document is ready for download.")

        # Provide a download link
        st.download_button(label='Download Anonymized Document',
                           data=anonymized_stream,
                           file_name='anonymized_resume.docx',
                           mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        st.error(f"An error occurred: {e}")

# Streamlit UI
st.title("Resume Anonymizer")

# File uploader
uploaded_file = st.file_uploader("Upload a resume", type=["docx"])
if uploaded_file is not None:
    anonymize_document(uploaded_file)
