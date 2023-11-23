import docx
import spacy

# Load spaCy model
nlp = spacy.load("en_core_web_sm")

def anonymize_document(r"G:\Python VS Code\aEnviroments\ResumeNameRemover\ResumeDocs\Calin_Andras_CV.docx")
    doc = docx.Document(r"G:\Python VS Code\aEnviroments\ResumeNameRemover\ResumeDocs\Calin_Andras_CV.docx")
    anonymized_doc = docx.Document()

    for para in doc.paragraphs:
        anonymized_text = para.text
        processed_text = nlp(anonymized_text)
        for entity in processed_text.ents:
            if entity.label_ == "PERSON":
                anonymized_text = anonymized_text.replace(entity.text, "[ANONYMIZED]")
        
        anonymized_doc.add_paragraph(anonymized_text)

    anonymized_doc.save("anonymized_resume.docx")