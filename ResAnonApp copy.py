import os
import docx
import spacy

file_path = r"G:\Python VS Code\aEnviroments\ResumeNameRemover\Calin Andras CV.docx"
# Check if the file exists
if not os.path.isfile(file_path):
    print(f"The file does not exist: {file_path}")
else:
    print(f"The file exists: {file_path}")
    # Continue with opening and processing the file

def anonymize_document(file_path):
    try:
        print(f"Processing file: {file_path}")
        nlp = spacy.load("en_core_web_sm")
        doc = docx.Document(file_path)
        anonymized_doc = docx.Document()

        for para in doc.paragraphs:
            anonymized_text = para.text
            processed_text = nlp(anonymized_text)
            for entity in processed_text.ents:
                if entity.label_ == "PERSON":
                    anonymized_text = anonymized_text.replace(entity.text, "[ANONYMIZED]")
            
            anonymized_doc.add_paragraph(anonymized_text)

        anonymized_doc.save(r"G:\Python VS Code\aEnviroments\ResumeNameRemover\anonymized_resume.docx")
        print("Anonymized document saved.")
    except Exception as e:
        print(f"An error occurred: {e}")
        raise

# Call the function with the path to the file
file_path = r"G:\Python VS Code\aEnviroments\ResumeNameRemover\Calin Andras CV.docx"
anonymize_document(file_path)

